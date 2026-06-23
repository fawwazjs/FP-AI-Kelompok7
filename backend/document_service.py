# Lokalator Document Parsing & Reconstruction Engine
# Handles PDF (PyMuPDF), DOCX (python-docx), legacy DOC extraction, and TXT.

import re
import textwrap
from pathlib import Path

from .translator_service import translate_and_classify, get_word_count

try:
    import fitz  # PyMuPDF
    PYMUPDF_AVAILABLE = True
except ImportError:
    PYMUPDF_AVAILABLE = False

try:
    import docx
    PYTHON_DOCX_AVAILABLE = True
except ImportError:
    PYTHON_DOCX_AVAILABLE = False


def _politeness_summary(target_lang: str, level: str) -> str:
    if target_lang == "jv":
        return "Krama Alus" if level == "high" else "Ngoko Lugu"
    if target_lang == "mad":
        return "Engghi-Bhanten" if level == "high" else "Enja-Iya"
    return "Netral"


def _summary(output_path: str, target_lang: str, level: str, translated_text: str, source_text: str) -> dict:
    return {
        "words_translated": get_word_count(source_text),
        "politeness_summary": _politeness_summary(target_lang, level),
        "file_created": output_path,
        "translated_text": translated_text.strip(),
        "source_text": source_text,
    }


def _translate_block(text: str, source_lang: str, target_lang: str, level: str) -> str:
    """Translate a block of text using Gemini+RAG, falling back to rule-based."""
    if not text.strip():
        return ""
    # Try Gemini + RAG first for document translation
    try:
        from .gemini_service import translate_with_gemini
        gemini_result = translate_with_gemini(text, source_lang, target_lang, level)
        if gemini_result and gemini_result.strip():
            return gemini_result
    except Exception:
        pass
    # Fallback to rule-based
    result = translate_and_classify(text, source_lang, target_lang, level)
    return result["translatedText"]


def _translate_lines(text: str, source_lang: str, target_lang: str, level: str) -> str:
    # Memisahkan teks berdasarkan paragraf atau baris kosong untuk menjaga konteks utuh
    # dan mencegah pemanggilan API per-baris yang menyebabkan rate-limit/lemot.
    paragraphs = re.split(r'(\n\s*\n)', text)
    translated_chunks = []
    
    current_chunk = ""
    for p in paragraphs:
        if not p.strip():
            current_chunk += p
            continue
            
        if len(current_chunk) + len(p) > 2000:
            if current_chunk.strip():
                translated_chunks.append(_translate_block(current_chunk, source_lang, target_lang, level))
            current_chunk = p
        else:
            current_chunk += p
            
    if current_chunk.strip():
        translated_chunks.append(_translate_block(current_chunk, source_lang, target_lang, level))
    elif current_chunk:
        translated_chunks.append(current_chunk)
        
    return "".join(translated_chunks)


def _read_txt(input_path: str) -> str:
    for encoding in ("utf-8", "utf-16", "latin-1"):
        try:
            return Path(input_path).read_text(encoding=encoding, errors="ignore")
        except Exception:
            continue
    return Path(input_path).read_bytes().decode("utf-8", errors="ignore")


def _write_translated_pdf(output_path: str, translated_text: str):
    if not PYMUPDF_AVAILABLE:
        Path(output_path).write_text(translated_text, encoding="utf-8")
        return

    doc = fitz.open()
    page = doc.new_page(width=595, height=842)  # A4 points
    margin_x = 48
    y = 56
    line_height = 14
    for paragraph in translated_text.split("\n"):
        wrapped = textwrap.wrap(paragraph, width=88) or [""]
        for line in wrapped:
            if y > 790:
                page = doc.new_page(width=595, height=842)
                y = 56
            page.insert_text((margin_x, y), line, fontsize=10)
            y += line_height
        y += 6

    doc.save(output_path)
    doc.close()


def process_and_translate_pdf(
    input_path: str,
    output_path: str,
    target_lang: str,
    source_lang: str = "id",
    level: str = "high"
) -> dict:
    if not PYMUPDF_AVAILABLE:
        source_text = Path(input_path).read_bytes().decode("utf-8", errors="ignore")
    else:
        doc = fitz.open(input_path)
        pages = [page.get_text() for page in doc]
        doc.close()
        source_text = "\n".join(pages)

    translated_text = _translate_lines(source_text, source_lang, target_lang, level)
    _write_translated_pdf(output_path, translated_text)
    return _summary(output_path, target_lang, level, translated_text, source_text)


def process_and_translate_docx(
    input_path: str,
    output_path: str,
    target_lang: str,
    source_lang: str = "id",
    level: str = "high"
) -> dict:
    if not PYTHON_DOCX_AVAILABLE:
        source_text = _read_txt(input_path)
        translated_text = _translate_lines(source_text, source_lang, target_lang, level)
        Path(output_path).write_text(translated_text, encoding="utf-8")
        return _summary(output_path, target_lang, level, translated_text, source_text)

    source_doc = docx.Document(input_path)
    valid_paragraphs = []

    for paragraph in source_doc.paragraphs:
        if paragraph.text.strip():
            valid_paragraphs.append(paragraph)

    for table in source_doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if paragraph.text.strip():
                        valid_paragraphs.append(paragraph)

    source_text_list = [p.text for p in valid_paragraphs]
    source_text_full = "\n".join(source_text_list)
    
    # Terjemahkan sekaligus menggunakan chunking untuk menghindari rate-limit
    translated_text_full = _translate_lines(source_text_full, source_lang, target_lang, level)
    translated_text_list = translated_text_full.split("\n")

    # Map teks yang diterjemahkan kembali ke run paragraf asli untuk mempertahankan format (font, warna)
    for i, paragraph in enumerate(valid_paragraphs):
        translated_text = translated_text_list[i] if i < len(translated_text_list) else paragraph.text
        if paragraph.runs:
            paragraph.runs[0].text = translated_text
            for j in range(1, len(paragraph.runs)):
                paragraph.runs[j].text = ""

    source_doc.save(output_path)
    return _summary(output_path, target_lang, level, translated_text_full, source_text_full)



def process_and_translate_txt(
    input_path: str,
    output_path: str,
    target_lang: str,
    source_lang: str = "id",
    level: str = "high"
) -> dict:
    source_text = _read_txt(input_path)
    translated_text = _translate_lines(source_text, source_lang, target_lang, level)
    Path(output_path).write_text(translated_text, encoding="utf-8")
    return _summary(output_path, target_lang, level, translated_text, source_text)


def process_and_translate_doc(
    input_path: str,
    output_path: str,
    target_lang: str,
    source_lang: str = "id",
    level: str = "high"
) -> dict:
    if PYTHON_DOCX_AVAILABLE:
        try:
            return process_and_translate_docx(input_path, output_path, target_lang, source_lang, level)
        except Exception:
            pass

    content = Path(input_path).read_bytes()

    utf16_text = ""
    try:
        raw_utf16 = content.decode("utf-16-le", errors="ignore")
        utf16_text = "".join(c for c in raw_utf16 if c.isprintable() or c in "\n\r\t")
    except Exception:
        pass

    ansi_text = "".join(chr(b) if (32 <= b <= 126 or b in (10, 13, 9)) else " " for b in content)
    raw_text = utf16_text if len(utf16_text.split()) > len(ansi_text.split()) else ansi_text

    clean_lines = []
    for line in raw_text.splitlines():
        cleaned = re.sub(r"[^a-zA-Z0-9\s.,!?()'\"\-]", "", line).strip()
        if len(cleaned) > 3:
            clean_lines.append(cleaned)

    source_text = "\n".join(clean_lines)
    translated_text = _translate_lines(source_text, source_lang, target_lang, level)
    Path(output_path).write_text(translated_text, encoding="utf-8")
    return _summary(output_path, target_lang, level, translated_text, source_text)
