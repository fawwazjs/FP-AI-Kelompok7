import tempfile
import unittest
from pathlib import Path

from backend.document_service import (
    PYTHON_DOCX_AVAILABLE,
    process_and_translate_docx,
    process_and_translate_txt,
)
from backend.translator_service import detect_language_and_register, translate_and_classify

if PYTHON_DOCX_AVAILABLE:
    import docx


class TranslationPipelineTests(unittest.TestCase):
    def test_supported_language_translation_cases(self):
        cases = [
            ("Saya ingin makan nasi goreng.", "id", "jv", "high", ["Kula", "dhahar", "sekul"]),
            ("saya ingin makan nasi goreng", "id", "jv", "low", ["Aku", "mangan", "sego"]),
            ("kamu mau pergi ke mana?", "id", "jv", "high", ["Panjenengan", "tindak", "pundi"]),
            ("terima kasih banyak atas bantuannya", "id", "jv", "high", ["Matur", "nuwun", "sanget"]),
            ("selamat pagi bagaimana kabar anda", "id", "jv", "low", ["Sugeng", "piye"]),
            ("nama saya ahmad saya tinggal di surabaya", "id", "jv", "high", ["Nami", "kula", "Surabaya"]),
            ("Saya makan nasi dan minum air.", "id", "jv", "high", ["Kula", "dhahar", "ngunjuk"]),
            ("Saya tidak bisa tidur sekarang.", "id", "jv", "low", ["Aku", "ora", "turu"]),
            ("rumah besar dan air sedikit", "id", "jv", "high", ["griya", "ageng", "toya"]),
            ("Merah besar", "id", "jv", "high", ["Abrit", "ageng"]),
            ("Saya ingin makan nasi goreng.", "id", "mad", "high", ["Bhiula", "neddha", "nase"]),
            ("saya ingin makan nasi goreng", "id", "mad", "low", ["Sengko", "ngakana", "nase"]),
            ("kamu mau pergi ke mana?", "id", "mad", "high", ["Panjhenengngan", "alomampaha"]),
            ("terima kasih banyak atas bantuannya", "id", "mad", "low", ["Sakalangkong", "raje"]),
            ("selamat pagi bagaimana kabar anda", "id", "mad", "low", ["de'remmah", "kabarra"]),
            ("nama saya ahmad saya tinggal di surabaya", "id", "mad", "high", ["bhiula", "Surabaya"]),
            ("Saya makan nasi dan minum air.", "id", "mad", "high", ["Bhiula", "neddha", "ngonjhung"]),
            ("Saya tidak bisa tidur sekarang.", "id", "mad", "low", ["Sengko", "tedhung", "sateya"]),
            ("ayah", "id", "mad", "high", ["ba"]),
            ("Aku mangan sego.", "jv", "id", "high", ["Saya", "makan", "nasi"]),
            ("Kula badhe dhahar sekul.", "jv", "id", "high", ["Saya", "ingin", "makan"]),
            ("Kowe arep lunga nang endi?", "jv", "id", "low", ["Kamu", "mau", "pergi"]),
            ("Sengko' terro ngakan nase'.", "mad", "id", "low", ["Saya", "ingin", "makan"]),
            ("Bhiula terro neddha nase'.", "mad", "id", "high", ["Saya", "ingin", "makan"]),
            ("Panjhenengngan badhi alomampaha ka dhimma?", "mad", "id", "high", ["Kamu", "mau", "pergi"]),
            ("Aku mangan sego.", "jv", "mad", "low", ["Sengko", "ngakan", "nase"]),
            ("Sengko' terro ngakan nase'.", "mad", "jv", "high", ["Kula", "dhahar", "sekul"]),
        ]
        self.assertGreaterEqual(len(cases), 24)

        for text, source, target, level, expected_parts in cases:
            with self.subTest(text=text, source=source, target=target, level=level):
                result = translate_and_classify(text, source, target, level)
                translated = result["translatedText"]
                self.assertTrue(translated.strip())
                for expected in expected_parts:
                    self.assertIn(expected, translated)

    def test_language_register_detection(self):
        cases = [
            ("Saya sedang makan nasi.", "Indonesia", "formal"),
            ("Gue pengen makan nasi.", "Indonesia", "informal"),
            ("Aku arep mangan sego.", "Jawa", "ngoko lugu"),
            ("Kula badhe dhahar sekul.", "Jawa", "krama alus"),
            ("Sengko' terro ngakan nase'.", "Madura", "Enja-Iya"),
            ("Panjhenengngan badhi alomampaha ka dhimma?", "Madura", "Engghi-bhunten"),
        ]

        for text, language, register in cases:
            with self.subTest(text=text):
                result = detect_language_and_register(text)
                self.assertEqual(result["language"], language)
                self.assertEqual(result["register"], register)

    def test_txt_document_translation_contains_actual_text(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            input_path = Path(temp_dir) / "sample.txt"
            output_path = Path(temp_dir) / "sample_translated.txt"
            input_path.write_text("Saya ingin makan nasi goreng.\nKamu mau pergi ke mana?", encoding="utf-8")

            summary = process_and_translate_txt(str(input_path), str(output_path), "jv", "id", "high")

            output = output_path.read_text(encoding="utf-8")
            self.assertIn("Kula", output)
            self.assertIn("Panjenengan", output)
            self.assertIn("Kula", summary["translated_text"])

    @unittest.skipUnless(PYTHON_DOCX_AVAILABLE, "python-docx is not installed")
    def test_docx_document_translation_contains_actual_text(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            input_path = Path(temp_dir) / "sample.docx"
            output_path = Path(temp_dir) / "sample_translated.docx"

            source_doc = docx.Document()
            source_doc.add_paragraph("Saya ingin makan nasi goreng.")
            source_doc.add_paragraph("Kamu mau pergi ke mana?")
            source_doc.save(input_path)

            summary = process_and_translate_docx(str(input_path), str(output_path), "jv", "id", "high")
            translated_doc = docx.Document(output_path)
            translated_text = "\n".join(paragraph.text for paragraph in translated_doc.paragraphs)

            self.assertIn("Kula", translated_text)
            self.assertIn("Panjenengan", translated_text)
            self.assertIn("Kula", summary["translated_text"])


if __name__ == "__main__":
    unittest.main()
